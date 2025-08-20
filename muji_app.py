# complete Working Code — NexaSight – AI Business Analyst (no theme toggle, updated hero styling)

import os
import io
import socket
import tempfile
import hashlib
import shutil
import streamlit as st
import pandas as pd

# --- PDF readers (graceful fallback) ---
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except Exception:
    PDFPLUMBER_AVAILABLE = False
    try:
        from pypdf import PdfReader
        PYPDF_AVAILABLE = True
    except Exception:
        PYPDF_AVAILABLE = False

from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings

# Optional exports
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

import matplotlib.pyplot as plt

# -----------------------------
# CONFIG
# -----------------------------
APP_NAME = "NexaSight – AI Business Analyst"
APP_TAGLINE = "Analyze multiple PDF and CSV/Excel files with AI-powered insights. Per-file summaries 📄 • multilingual 🌍 • export reports • charts 📈"

st.set_page_config(
    page_title=APP_NAME,
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# (LAN IP helper kept but no UI shown)
def _lan_ip() -> str:
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.settimeout(0.2)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
    except Exception:
        ip = "127.0.0.1"
    finally:
        try:
            s.close()
        except Exception:
            pass
    return ip

LAN_IP = _lan_ip()

# -----------------------------
# Branding (single professional theme, no toggle)
# -----------------------------
def palette():
    """
    Professional, calm palette optimized for dark UI without a runtime toggle.
    """
    return dict(
        bg="#0A0F1C",      # deep navy background
        card="#0E1528",    # slightly lighter card
        text="#E8F0FB",    # near-white text
        sub="#A9B6D3",     # muted subtitle
        border="#1E2A42",  # cool border
        grad_a="#1B2A4A",  # hero gradient start (navy)
        grad_b="#233B6E",  # hero gradient mid (indigo)
        grad_c="#2C6FB3",  # hero gradient end (blue)
        accent="#3BB9E3",  # cyan-accent
        warn="#F59E0B",
    )

def make_logo_svg():
    c = palette()
    return f"""
<svg width="220" height="48" viewBox="0 0 220 48" xmlns="http://www.w3.org/2000/svg" role="img" aria-label="NexaSight logo">
  <defs>
    <linearGradient id="g1" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0%" stop-color="{c['grad_a']}"/>
      <stop offset="60%" stop-color="{c['grad_b']}"/>
      <stop offset="100%" stop-color="{c['grad_c']}"/>
    </linearGradient>
  </defs>
  <!-- Hex prism -->
  <g transform="translate(6,6)">
    <polygon points="18,0 36,10 36,28 18,38 0,28 0,10" fill="url(#g1)"/>
    <!-- Monogram N -->
    <path d="M7 27 L7 11 L13 11 L26 24 L26 11 L29 11 L29 27 L23 27 L10 14 L10 27 Z"
          fill="#ffffff" opacity="0.95"/>
  </g>

  <!-- Wordmark -->
  <g transform="translate(60,33)">
    <text x="0" y="0" font-family="Inter, ui-sans-serif, -apple-system, Segoe UI, Roboto"
          font-weight="800" font-size="22" fill="{c['text']}">NexaSight</text>
  </g>
</svg>
"""

# Global CSS (no brittle class names)
def inject_css():
    c = palette()
    st.markdown(
        f"""
<style>
:root {{
  --bg: {c['bg']};
  --card: {c['card']};
  --text: {c['text']};
  --sub: {c['sub']};
  --border: {c['border']};
  --accent: {c['accent']};
  --warn: {c['warn']};
}}
html, body, .stApp {{ background: var(--bg) !important; color: var(--text) !important; }}

/* HERO (updated color + professional sizing) */
.hero {{
  background: linear-gradient(90deg, {c['grad_a']}, {c['grad_b']} 45%, {c['grad_c']});
  padding: 24px 26px; border-radius: 18px; box-shadow: 0 10px 28px rgba(2, 6, 23, .35);
  color: #fff; margin-bottom: 14px;
}}
.hero .row {{ display:flex; align-items:center; gap:18px; flex-wrap: wrap; }}
.hero .title {{
  font-size: 24px; font-weight: 800; margin: 0;
}}
@media (max-width: 640px) {{
  .hero .title {{ font-size: 22px; }}
}}
.hero .tagline {{ font-size: 13px; opacity: .85; margin-top: 4px; }}

/* CONTENT CARDS */
.block {{
  background: var(--card); border: 1px solid var(--border);
  border-radius: 16px; padding: 14px 16px; margin: 8px 0;
}}

/* Sidebar tweaks */
section[data-testid="stSidebar"] > div {{ background: var(--card) }}

/* Upload callout */
.callout {{
  border:1px dashed rgba(255,255,255,.35); border-radius:12px; padding:12px 14px; margin-top:8px;
  background: rgba(255,255,255,.04);
}}

/* Chat bubbles */
.chat-message {{ padding: .85rem; margin: .45rem 0; border-radius: 12px; font-size: 0.98rem; line-height: 1.5; }}
.chat-message.user {{ background: rgba(59,185,227,.15); border:1px solid rgba(59,185,227,.35); color: var(--text); }}
.chat-message.bot {{ background: rgba(148,163,184,.12); border:1px solid var(--border); color: var(--text); }}

/* Summary bullet block */
.summary-block {{ padding:.6rem .9rem; border-radius:10px; border-left:5px solid var(--warn);
                 margin-bottom:.5rem; background: rgba(252,211,77,.08); color: var(--text); }}
.summary-title {{ font-weight: 700; margin-bottom: .25rem; }}
ul {{ margin: 0 0 .25rem 1rem; padding: 0; }}
.logo-wrap {{ margin: 0; }}

/* Buttons (subtle lift) */
button[kind="secondary"]{{ box-shadow:none }}
</style>
""",
        unsafe_allow_html=True,
    )

# -----------------------------
# HEADER (Logo + Hero Title Block)
# -----------------------------
inject_css()

st.markdown(
    f"""
<div class="hero">
  <div class="row">
    <div class="logo-wrap">{make_logo_svg()}</div>
    <div>
      <p class="title">{APP_NAME}</p>
      <p class="tagline">{APP_TAGLINE}</p>
    </div>
  </div>
</div>
<div class="callout">📌 Please upload one or more files to start.</div>
""",
    unsafe_allow_html=True,
)

# -----------------------------
# API Key (unchanged model usage)
# -----------------------------
os.environ.setdefault("GOOGLE_API_KEY", "AIzaSyD-Ct8OiZKvGHHFEnux8i0KkQdigDIZ6XA")

# -----------------------------
# Sidebar - Multiple File Upload + Language
# -----------------------------
st.sidebar.header("📂 Upload Files")
uploaded_files = st.sidebar.file_uploader(
    "Upload PDF, CSV, or Excel files",
    type=["pdf", "csv", "xlsx"],
    accept_multiple_files=True,
)

language = st.sidebar.selectbox(
    "🌍 Summary & Chat Language",
    ["English", "Spanish", "French", "Hindi", "German", "Chinese", "Arabic", "Urdu"],
    index=0,
)

with st.sidebar.expander("⚙️ Cache, Refresh & Export Options", expanded=False):
    clear_vec_cache = st.button("🧹 Clear Vector Store Cache")
    auto_refresh = st.checkbox(
        "Auto-refresh when uploads change",
        value=True,
        help="Automatically rebuild index and summaries whenever the uploaded file set changes.",
    )
    if st.button("🔄 Refresh App Now"):
        # Manual refresh: reset state and rerun
        st.session_state.file_summaries = None
        st.session_state.summary_text = None
        st.session_state.analytics = {}
        st.session_state.current_cache_key = None
        st.cache_data.clear()
        st.cache_resource.clear()
        st.toast("App refreshed.", icon="🔄")
        st.rerun()

    st.caption("If files change but keep the same names, clear or refresh to rebuild the index.")
    st.write("Export options:")
    enable_pdf_export = st.checkbox("Enable PDF export (reportlab)", value=REPORTLAB_AVAILABLE)
    enable_docx_export = st.checkbox("Enable Word export (python-docx)", value=DOCX_AVAILABLE)

# -----------------------------
# Session-state initialization
# -----------------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "file_summaries" not in st.session_state:
    st.session_state.file_summaries = None
if "summary_text" not in st.session_state:
    st.session_state.summary_text = None
if "analytics" not in st.session_state:
    st.session_state.analytics = {}
# Track current cache key to detect upload changes
if "current_cache_key" not in st.session_state:
    st.session_state.current_cache_key = None

# -----------------------------
# Helpers
# -----------------------------
def _hash_bytes(data: bytes) -> str:
    import hashlib as _h
    return _h.sha256(data).hexdigest()

def _file_fingerprint(file_bytes: bytes, filename: str) -> str:
    base = f"{filename}|{len(file_bytes)}|{_hash_bytes(file_bytes)}"
    return hashlib.sha256(base.encode("utf-8")).hexdigest()

@st.cache_resource(show_spinner=False)
def get_embeddings():
    return GoogleGenerativeAIEmbeddings(model="models/embedding-001")

@st.cache_resource(show_spinner=False)
def get_llm():
    # MODEL UNCHANGED
    return ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.2)

@st.cache_data(show_spinner=False)
def extract_docs_from_file(file_bytes: bytes, filename: str):
    """
    Return list of text chunks extracted from a file.
    Robust: uses pdfplumber if installed; else falls back to pypdf; else emits a marker.
    """
    docs = []
    ext = filename.split(".")[-1].lower()
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    try:
        if ext == "pdf":
            if PDFPLUMBER_AVAILABLE:
                import io as _io
                with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                    def process_pdf_page(page):
                        try:
                            text = page.extract_text() or ""
                            return splitter.split_text(text) if text.strip() else []
                        except Exception as e:
                            return [f"[PDF_PAGE_READ_ERROR:{filename}] {e}"]
                    with ThreadPoolExecutor() as executor:
                        for page_docs in executor.map(process_pdf_page, pdf.pages):
                            docs.extend(page_docs)
            elif PYPDF_AVAILABLE:
                text_all = []
                try:
                    from io import BytesIO
                    reader = PdfReader(BytesIO(file_bytes))
                    for p in reader.pages:
                        try:
                            text_all.append(p.extract_text() or "")
                        except Exception as e:
                            text_all.append(f"[PDF_PAGE_READ_ERROR:{filename}] {e}")
                    docs.extend(splitter.split_text("\n".join(text_all)))
                except Exception as e:
                    docs.append(f"[READ_ERROR:{filename}] {e}")
            else:
                docs.append(f"[MISSING_PDF_BACKEND:{filename}] Install 'pdfplumber' or 'pypdf'.")
        elif ext in ["csv", "xlsx"]:
            try:
                if ext == "csv":
                    import io as _io
                    df = pd.read_csv(_io.BytesIO(file_bytes), on_bad_lines="skip", engine="python")
                else:
                    from io import BytesIO
                    df = pd.read_excel(BytesIO(file_bytes))
                csv_text = df.to_csv(index=False)
                docs.extend(splitter.split_text(csv_text))
            except Exception as e:
                docs.append(f"[READ_ERROR:{filename}] {e}")
        else:
            docs.append(f"[UNSUPPORTED:{filename}]")
    except Exception as e:
        docs.append(f"[READ_ERROR:{filename}] {e}")
    return docs

def vector_cache_dir(cache_key: str) -> str:
    root = os.path.join(tempfile.gettempdir(), "nexasight_cache")
    os.makedirs(root, exist_ok=True)
    path = os.path.join(root, cache_key)
    os.makedirs(path, exist_ok=True)
    return path

def vectorstore_exists(path: str) -> bool:
    return os.path.exists(os.path.join(path, "index.faiss")) and os.path.exists(os.path.join(path, "index.pkl"))

# -----------------------------
# Export helpers (PDF/DOCX/TXT)
# -----------------------------
def _generate_pdf_bytes(title: str, summaries: dict) -> bytes:
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab not available")
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 40
    y = height - margin
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, title)
    y -= 30
    c.setFont("Helvetica", 11)
    for fname, summary in summaries.items():
        if y < 100:
            c.showPage()
            y = height - margin
            c.setFont("Helvetica", 11)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(margin, y, f"{fname}:")
        y -= 18
        c.setFont("Helvetica", 10)
        for line in summary.splitlines():
            for chunk in [line[i:i+90] for i in range(0, len(line), 90)]:
                if y < 60:
                    c.showPage()
                    y = height - margin
                c.drawString(margin+10, y, chunk)
                y -= 14
        y -= 8
    c.save()
    buffer.seek(0)
    return buffer.read()

def _generate_docx_bytes(title: str, summaries: dict) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not available")
    doc = docx.Document()
    doc.add_heading(title, level=1)
    for fname, summary in summaries.items():
        doc.add_heading(fname, level=2)
        for line in summary.splitlines():
            doc.add_paragraph(line)
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()

def _generate_text_bytes(title: str, summaries: dict) -> bytes:
    buf = io.StringIO()
    buf.write(title + "\n\n")
    for fname, summary in summaries.items():
        buf.write(fname + ":\n")
        buf.write(summary + "\n\n")
    return buf.getvalue().encode("utf-8")

# -----------------------------
# Build / Load Vector Store
# -----------------------------
documents, cache_hit, cache_key, vecstore_path = [], False, None, None
vectorstore = None

if uploaded_files:
    file_bytes_list = [(f.getvalue(), f.name) for f in uploaded_files]
    fingerprints = [_file_fingerprint(b, name) for b, name in file_bytes_list]
    fingerprints.sort()
    cache_key = hashlib.sha256("|".join(fingerprints).encode("utf-8")).hexdigest()

    # Detect change in uploaded file set and auto-refresh state
    if st.session_state.current_cache_key != cache_key:
        # auto_refresh is defined above inside the expander; fall back to True if not found
        if 'auto_refresh' not in locals():
            auto_refresh = True
        if auto_refresh:
            st.session_state.file_summaries = None
            st.session_state.summary_text = None
            st.session_state.analytics = {}
            st.toast("Detected new uploads → rebuilding index & summaries…", icon="🔄")
        st.session_state.current_cache_key = cache_key

    vecstore_path = vector_cache_dir(cache_key)

    if clear_vec_cache and os.path.isdir(vecstore_path):
        # Clear only the current key's FAISS folder
        for fname in ["index.faiss", "index.pkl"]:
            fpath = os.path.join(vecstore_path, fname)
            if os.path.exists(fpath):
                os.remove(fpath)
        st.session_state.file_summaries = None
        st.session_state.summary_text = None
        st.toast("Vector store cache cleared for current upload set.", icon="🧹")

    embeddings = get_embeddings()

    if vectorstore_exists(vecstore_path):
        try:
            vectorstore = FAISS.load_local(
                vecstore_path, embeddings, allow_dangerous_deserialization=True
            )
            cache_hit = True
            st.info("⚡ Using cached index (no re-embedding).")
        except Exception as e:
            st.warning(f"Failed to load cached index, will rebuild. ({e})")
            vectorstore = None

    if vectorstore is None:
        st.info("📄 Processing uploaded files and building index…")
        with ThreadPoolExecutor() as executor:
            for file_docs in executor.map(
                lambda pair: extract_docs_from_file(pair[0], pair[1]), file_bytes_list
            ):
                documents.extend(file_docs)
        if documents:
            try:
                vectorstore = FAISS.from_texts(documents, embeddings)
                vectorstore.save_local(vecstore_path)
                st.success("✅ Index built and cached for future runs.")
            except Exception as e:
                st.error(f"Failed to build vectorstore: {e}")

    # Auto Summary Generation
    if vectorstore is not None:
        if st.session_state.file_summaries is None:
            st.session_state.file_summaries = {}
            llm = get_llm()
            splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=150)

            for file_obj in uploaded_files:
                file_docs = extract_docs_from_file(file_obj.getvalue(), file_obj.name)
                file_text = "\n".join(file_docs).strip()
                if not file_text:
                    st.session_state.file_summaries[file_obj.name] = "[No readable content]"
                    continue

                file_chunks = splitter.split_text(file_text)
                partial_summaries = []
                for chunk in file_chunks:
                    prompt = (
                        f"You are summarizing content from the file '{file_obj.name}'. "
                        f"Write the summary **in {language} only**, using 5 short bullet points max. "
                        f"Remove unnecessary spaces and blank lines. Be concise:\n\n{chunk}"
                    )
                    try:
                        result = llm.invoke(prompt)
                        partial_summaries.append(
                            result.content if hasattr(result, "content") else str(result)
                        )
                    except Exception as e:
                        partial_summaries.append(f"[LLM_ERROR:{e}]")

                combine_prompt = (
                    f"Combine these partial summaries into a single clean bullet list for '{file_obj.name}' in {language}. "
                    f"Max 5 bullets, no blank lines, no repetition:\n\n" + "\n".join(partial_summaries)
                )
                try:
                    final_result = llm.invoke(combine_prompt)
                    clean_summary = final_result.content.strip()
                except Exception as e:
                    clean_summary = "[SUMMARY_COMBINE_ERROR] " + str(e)
                st.session_state.file_summaries[file_obj.name] = clean_summary

            combined_text = "\n".join(
                f"<div class='summary-block'><div class='summary-title'>{fname}:</div><ul>{''.join(f'<li>{line.strip()}</li>' for line in summary.splitlines() if line.strip())}</ul></div>"
                for fname, summary in st.session_state.file_summaries.items()
            )
            st.session_state.summary_text = combined_text

            # Also prepare basic analytics for CSV/XLSX files
            st.session_state.analytics = {}
            for file_obj in uploaded_files:
                name = file_obj.name
                ext = name.split('.')[-1].lower()
                if ext in ['csv', 'xlsx']:
                    try:
                        if ext == 'csv':
                            df = pd.read_csv(io.BytesIO(file_obj.getvalue()), on_bad_lines='skip', engine='python')
                        else:
                            df = pd.read_excel(io.BytesIO(file_obj.getvalue()))
                        desc = df.describe(include='all').transpose()
                        st.session_state.analytics[name] = {'df': df, 'describe': desc}
                    except Exception as e:
                        st.session_state.analytics[name] = {'error': str(e)}

# -----------------------------
# Chat + UI
# -----------------------------
if uploaded_files and (vectorstore is not None):
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    llm = get_llm()
    qa_chain = RetrievalQA.from_chain_type(llm, retriever=retriever)

    with st.expander("📦 Loaded Files", expanded=False):
        for f in uploaded_files:
            st.write(f"- {f.name} ({len(f.getvalue())/1024:.1f} KB)")
        st.write(f"🧠 Cache Key: `{cache_key[:12]}...`")
        st.write("Cache Status: " + ("**HIT** ⚡" if cache_hit else "**MISS** 🧱"))

    col1, col2, col3 = st.columns([0.33, 0.33, 0.34])
    with col1:
        if st.button("🧹 Clear Chat"):
            st.session_state.chat_history = []
            st.success("Chat cleared.")
    with col2:
        if st.button("💾 Save Chat"):
            if st.session_state.chat_history:
                chat_text = "\n".join([f"{role.upper()}: {msg}" for role, msg in st.session_state.chat_history])
                st.download_button(
                    label="📥 Download Chat",
                    data=chat_text,
                    file_name=f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                )
            else:
                st.warning("No chat history to save.")
    with col3:
        if st.button("🔄 Refresh App"):
            st.session_state.file_summaries = None
            st.session_state.summary_text = None
            st.session_state.analytics = {}
            st.toast("Refresh requested. Rebuilding…", icon="🔄")
            st.rerun()

    left_col, right_col = st.columns([0.35, 0.65])
    with left_col:
        if st.session_state.summary_text:
            with st.expander("📄 Auto Summary (per file)", expanded=False):
                st.markdown(st.session_state.summary_text, unsafe_allow_html=True)

        # Export Report UI
        with st.expander("📤 Export Summaries / Reports", expanded=False):
            st.write("Download a combined report of all file summaries")
            report_title = st.text_input(
                "Report title", value=f"{APP_NAME} Summary {datetime.now().strftime('%Y-%m-%d')}"
            )
            col_a, col_b = st.columns([0.5, 0.5])
            with col_a:
                if st.button("Export as PDF"):
                    try:
                        if REPORTLAB_AVAILABLE and enable_pdf_export:
                            pdf_bytes = _generate_pdf_bytes(
                                report_title, st.session_state.file_summaries or {}
                            )
                            st.download_button(
                                "Download PDF",
                                data=pdf_bytes,
                                file_name=f"{report_title}.pdf",
                                mime='application/pdf',
                            )
                        else:
                            txt_bytes = _generate_text_bytes(
                                report_title, st.session_state.file_summaries or {}
                            )
                            st.download_button(
                                "Download TXT (reportlab not installed)",
                                data=txt_bytes,
                                file_name=f"{report_title}.txt",
                                mime='text/plain',
                            )
                    except Exception as e:
                        st.error(f"Failed to create PDF: {e}")
            with col_b:
                if st.button("Export as Word"):
                    try:
                        if DOCX_AVAILABLE and enable_docx_export:
                            docx_bytes = _generate_docx_bytes(
                                report_title, st.session_state.file_summaries or {}
                            )
                            st.download_button(
                                "Download DOCX",
                                data=docx_bytes,
                                file_name=f"{report_title}.docx",
                                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            )
                        else:
                            txt_bytes = _generate_text_bytes(
                                report_title, st.session_state.file_summaries or {}
                            )
                            st.download_button(
                                "Download TXT (python-docx not installed)",
                                data=txt_bytes,
                                file_name=f"{report_title}.txt",
                                mime='text/plain',
                            )
                    except Exception as e:
                        st.error(f"Failed to create DOCX: {e}")

    with right_col:
        # Analytics & Charts UI for each CSV/XLSX
        with st.expander("📊 Charts & Analytics", expanded=False):
            if st.session_state.analytics:
                for fname, info in st.session_state.analytics.items():
                    st.markdown(f"**{fname}**")
                    if 'error' in info:
                        st.error(f"Failed to read {fname}: {info['error']}")
                        continue
                    df = info['df']
                    desc = info['describe']
                    st.write("Preview of data (first 5 rows):")
                    st.dataframe(df.head())

                    st.write("Descriptive statistics:")
                    st.dataframe(desc)

                    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
                    if numeric_cols:
                        st.write("Numeric columns detected:", numeric_cols)
                        chosen = st.selectbox(
                            f"Choose column to visualize ({fname})",
                            options=numeric_cols,
                            key=f"vis_{fname}",
                        )

                        if chosen:
                            fig, ax = plt.subplots()
                            try:
                                ax.hist(df[chosen].dropna().values)
                                ax.set_title(f"Histogram: {chosen}")
                                ax.set_xlabel(chosen)
                                ax.set_ylabel('Count')
                                st.pyplot(fig)
                            except Exception as e:
                                st.error(f"Failed to draw histogram for {chosen}: {e}")

                            if len(df) > 1:
                                fig2, ax2 = plt.subplots()
                                try:
                                    ax2.plot(df[chosen].dropna().values)
                                    ax2.set_title(f"Line plot (index vs {chosen})")
                                    ax2.set_xlabel('Index')
                                    ax2.set_ylabel(chosen)
                                    st.pyplot(fig2)
                                except Exception as e:
                                    st.error(f"Failed to draw line plot for {chosen}: {e}")
                    else:
                        st.info("No numeric columns found for plotting.")
            else:
                st.info("No analytics available for uploaded files.")

        # Chat form
        with st.form(key="chat_form", clear_on_submit=True):
            query = st.text_input(
                f"💬 Ask a question about your data ({language}):",
                key="user_query",
                placeholder="Type your question here...",
            )
            submit_button = st.form_submit_button("Ask")

        if submit_button and query:
            with st.spinner("🤔 Thinking..."):
                prompt_for_chain = (
                    f"Answer the following question in {language} only, without switching to any other language: {query}"
                )
                try:
                    result = qa_chain.run(prompt_for_chain)
                except Exception as e:
                    result = f"[Error running QA chain: {e}]"

            st.session_state.chat_history.insert(0, ("bot", result))
            st.session_state.chat_history.insert(0, ("user", query))

        for role, message in st.session_state.chat_history:
            if role == "user":
                st.markdown(f"<div class='chat-message user'>{message}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='chat-message bot'>{message}</div>", unsafe_allow_html=True)

elif uploaded_files:
    st.info("📌 Upload processed, but no content was indexed.")
else:
    pass  # initial callout already shown in hero

# -----------------------------
# Footer
# -----------------------------
st.markdown("---")
st.markdown(
    f"💡 **{APP_NAME}** — AI-powered insights from multiple business files (per-file summaries • multi-language • export • charts)."
)

# Helpful hint if optional libraries missing
if not REPORTLAB_AVAILABLE:
    st.info("Tip: Install 'reportlab' to enable nicer PDF exports: pip install reportlab")
if not DOCX_AVAILABLE:
    st.info("Tip: Install 'python-docx' to enable Word exports: pip install python-docx")
if not REPORTLAB_AVAILABLE and not DOCX_AVAILABLE:
    pass
if not PDFPLUMBER_AVAILABLE and 'PYPDF_AVAILABLE' in globals() and not PYPDF_AVAILABLE:
    st.info("Tip: Install 'pdfplumber' or 'pypdf' to read PDFs: pip install pdfplumber pypdf")







